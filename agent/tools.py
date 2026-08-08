"""Agent 工具集：执行命令、文件操作、系统信息。"""
import os
import shutil  # noqa: F401  保留用于未来扩展
import subprocess
import time
import psutil


class ToolError(Exception):
    """工具执行错误。"""


def run_command(cmd: str, timeout: int = 10) -> str:
    """执行 shell 命令，带超时保护。"""
    if not cmd.strip():
        raise ToolError("命令为空")
    try:
        proc = subprocess.run(
            cmd,
            shell=True,
            capture_output=True,
            text=True,
            timeout=timeout,
        )
        output = proc.stdout or (proc.stderr or "")
        code = proc.returncode
    except subprocess.TimeoutExpired:
        raise ToolError(f"命令执行超时（超过 {timeout} 秒）")
    except Exception as e:  # noqa: BLE001
        raise ToolError(f"命令执行失败：{e}")

    result = output.strip()
    if proc.returncode != 0:
        result = f"[退出码 {code}] {result}"
    return result or "(无输出)"


def list_dir(path: str = ".") -> str:
    """列出目录内容。"""
    try:
        entries = sorted(os.listdir(path))
    except OSError as e:
        raise ToolError(f"无法读取目录：{e}")
    if not entries:
        return "(空目录)"
    lines = []
    for name in entries:
        full = os.path.join(path, name)
        marker = "/" if os.path.isdir(full) else ""
        lines.append(f"{name}{marker}")
    return "\n".join(lines)


def read_file(path: str, max_bytes: int = 200_000) -> str:
    """读取文本文件内容（限制大小防止刷屏）。"""
    if not os.path.isfile(path):
        raise ToolError(f"文件不存在：{path}")
    size = os.path.getsize(path)
    if size > max_bytes:
        raise ToolError(f"文件过大（{size} 字节），超过上限 {max_bytes} 字节")
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    except OSError as e:
        raise ToolError(f"读取文件失败：{e}")


def open_path(path: str = ".") -> str:
    """用系统默认程序打开路径。"""
    path = os.path.abspath(path)
    if not os.path.exists(path):
        raise ToolError(f"路径不存在：{path}")
    try:
        if sys_platform_is_mac():
            subprocess.Popen(["open", path])
        elif sys_platform_is_win():
            os.startfile(path)  # type: ignore[attr-defined]
        else:
            subprocess.Popen(["xdg-open", path])
        return f"已打开：{path}"
    except OSError as e:
        raise ToolError(f"打开失败：{e}")


def make_dir(path: str) -> str:
    """创建目录（含父目录）。"""
    try:
        os.makedirs(path, exist_ok=True)
    except OSError as e:
        raise ToolError(f"创建目录失败：{e}")
    return f"已创建目录：{os.path.abspath(path)}"


def get_sysinfo() -> str:
    """获取系统信息。"""
    try:
        cpu_percent = psutil.cpu_percent(interval=0.3)
        mem = psutil.virtual_memory()
        disk = psutil.disk_usage("/")
        boot = psutil.boot_time()
        boot_str = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(boot))
        return (
            f"💻 系统信息\n"
            f"----------------------\n"
            f"操作系统: {os.uname().sysname} {os.uname().release}\n"
            f"CPU 使用率: {cpu_percent}%\n"
            f"内存: {mem.used / 1e9:.1f}G / {mem.total / 1e9:.1f}G ({mem.percent}%)\n"
            f"磁盘: {disk.used / 1e9:.1f}G / {disk.total / 1e9:.1f}G ({disk.percent}%)\n"
            f"开机时间: {boot_str}\n"
            f"当前用户: {os.getlogin()}"
        )
    except Exception as e:  # noqa: BLE001
        return f"获取系统信息失败：{e}"


def sys_platform_is_mac() -> bool:
    return os.name == "posix" and os.uname().sysname == "Darwin"


def sys_platform_is_win() -> bool:
    return os.name == "nt"


# ------------------------------------------------------------------
# 多格式文件读取（txt / csv / docx / xlsx / pdf）
# ------------------------------------------------------------------
SUPPORTED_EXTS = {".txt", ".md", ".log", ".csv", ".docx", ".xlsx", ".pdf"}


def read_file_content(path: str, max_bytes: int = 1_000_000) -> str:
    """按扩展名智能解析文件内容，返回适合传给大模型的文本。

    支持：.txt/.md/.log/.csv（直接文本）、.docx（Word）、.xlsx（Excel）、.pdf。
    若文件过大或二进制格式无法解析，抛出 ToolError。
    """
    if not os.path.isfile(path):
        raise ToolError(f"文件不存在：{path}")

    ext = os.path.splitext(path)[1].lower()
    size = os.path.getsize(path)
    if size > max_bytes:
        raise ToolError(f"文件过大（{size} 字节），超过上限 {max_bytes} 字节")

    try:
        if ext in (".txt", ".md", ".log", ".csv"):
            return _read_plain(path)
        if ext == ".docx":
            return _read_docx(path)
        if ext == ".xlsx":
            return _read_xlsx(path)
        if ext == ".pdf":
            return _read_pdf(path)
        raise ToolError(
            f"不支持的文件类型：{ext}。支持：{', '.join(sorted(SUPPORTED_EXTS))}"
        )
    except ToolError:
        raise
    except Exception as e:  # noqa: BLE001
        raise ToolError(f"解析文件失败：{e}")


def _read_plain(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _read_docx(path: str) -> str:
    from docx import Document

    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # 顺便提取表格内容
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    content = "\n".join(parts)
    return content or "(Word 文档无可见文本)"


def _read_xlsx(path: str) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    lines = []
    for ws in wb.worksheets:
        lines.append(f"[工作表] {ws.title}")
        for row in ws.iter_rows(values_only=True):
            vals = [
                (str(c).strip() if c is not None else "")
                for c in row
                if c is not None and str(c).strip()
            ]
            if vals:
                lines.append(" | ".join(vals))
    wb.close()
    return "\n".join(lines) if lines else "(Excel 无内容)"


def _read_pdf(path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text.strip())
    return "\n".join(pages) if pages else "(PDF 无可提取文本，可能是扫描件)"

